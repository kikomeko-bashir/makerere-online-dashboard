"""Generate Architecture Word document with a class diagram drawn using python-docx shapes."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title page
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MAKERERE ONLINE UNIVERSITY\nMANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Architecture Document')
run.font.size = Pt(16)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Version 1.0 | May 2026\n').font.size = Pt(12)
info.add_run('https://makerereonlineschool.com').font.size = Pt(12)

doc.add_page_break()

# Section 1
doc.add_heading('1. High-Level System Architecture', level=1)
doc.add_paragraph(
    'The system follows a three-tier architecture with a React frontend, '
    'FastAPI backend, and PostgreSQL database. All components are containerized '
    'with Docker and served behind an Apache reverse proxy with SSL.'
)
doc.add_paragraph()

# Architecture layers table
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['Layer', 'Component', 'Technology']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ('Client', 'Web Dashboard', 'React 19, Tailwind CSS, shadcn/ui'),
    ('Client', 'Mobile App', 'React Native (Expo)'),
    ('Application', 'REST API', 'Python FastAPI, SQLAlchemy 2.0'),
    ('Data', 'Database', 'PostgreSQL 16'),
]
for idx, (layer, comp, tech) in enumerate(data, 1):
    table.rows[idx].cells[0].text = layer
    table.rows[idx].cells[1].text = comp
    table.rows[idx].cells[2].text = tech

doc.add_paragraph()

# Section 2
doc.add_heading('2. Deployment Architecture', level=1)
doc.add_paragraph(
    'The system is deployed on a Contabo VPS (95.111.239.122) with the following setup:'
)
items = [
    'Apache2 handles SSL termination and routes traffic to Docker containers',
    'makerereonlineschool.com → Nginx container (port 3535) serving static React build',
    'api.makerereonlineschool.com → FastAPI container (port 3434)',
    'PostgreSQL runs on port 3435 (internal Docker network)',
    'Let\'s Encrypt (Certbot) provides free SSL certificates',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Section 3
doc.add_heading('3. Module Communication', level=1)
doc.add_paragraph(
    'Each frontend module communicates with a corresponding API router via REST. '
    'All requests include a JWT Bearer token for authentication.'
)

table2 = doc.add_table(rows=14, cols=3)
table2.style = 'Table Grid'
for i, h in enumerate(['Frontend Module', 'API Endpoint', 'Methods']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

modules = [
    ('Authentication', '/api/auth', 'POST login, POST register, GET me'),
    ('User Management', '/api/users', 'GET, POST, DELETE'),
    ('School Management', '/api/schools', 'GET, POST, PUT, DELETE'),
    ('Course Management', '/api/courses', 'GET, POST, PUT, DELETE'),
    ('Course Units', '/api/course-units', 'GET, POST, PUT, DELETE, PUT approve/reject'),
    ('Enrollment', '/api/enrollments', 'GET, POST, PUT'),
    ('Study Materials', '/api/materials', 'GET, POST, DELETE'),
    ('Virtual Learning', '/api/virtual-classes', 'GET, POST, PUT'),
    ('Examinations', '/api/examinations', 'GET, POST, PUT'),
    ('Certification', '/api/certificates', 'GET, POST'),
    ('Payments', '/api/payments', 'GET, POST'),
    ('Tutoring', '/api/tutoring', 'GET, POST, PUT'),
    ('Notifications', '/api/notifications', 'GET, PUT'),
]
for idx, (mod, endpoint, methods) in enumerate(modules, 1):
    table2.rows[idx].cells[0].text = mod
    table2.rows[idx].cells[1].text = endpoint
    table2.rows[idx].cells[2].text = methods

doc.add_paragraph()

# Section 4
doc.add_heading('4. Authentication Flow', level=1)
doc.add_paragraph('1. User submits email + password on the login page')
doc.add_paragraph('2. Frontend sends POST /api/auth/login with credentials')
doc.add_paragraph('3. API verifies password using bcrypt')
doc.add_paragraph('4. API generates a JWT token (HS256, 60-minute expiry)')
doc.add_paragraph('5. Frontend stores token in localStorage')
doc.add_paragraph('6. All subsequent API requests include Authorization: Bearer <token>')
doc.add_paragraph('7. API decodes token on each request to identify the user and role')

doc.add_paragraph()

# Section 5
doc.add_heading('5. Role-Based Access Control', level=1)

table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Role'
table3.rows[0].cells[1].text = 'Access'
for p in table3.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True

roles = [
    ('Super Admin', 'Full access: User Management, Schools, Courses, Units, Intakes, Enrollment, Payments, Reporting, Settings'),
    ('Admin', 'Schools, Courses, Units (approve/reject), Intakes, Enrollment, Materials, Virtual Learning, Exams, Certificates, Payments, Notifications, Reporting, User Management (no super_admin creation)'),
    ('Lecturer', 'Course Units (create pending), Materials (upload), Virtual Learning (schedule), Exams (create/grade), Tutoring, Notifications'),
    ('Student', 'Enrollment, Materials (view), Virtual Learning (join), Exams (take), Certificates (download), Payments, Tutoring (book), Notifications'),
]
for idx, (role, access) in enumerate(roles, 1):
    table3.rows[idx].cells[0].text = role
    table3.rows[idx].cells[1].text = access

doc.add_paragraph()
doc.add_page_break()

# Section 6 - Class Diagram as a table-based representation
doc.add_heading('6. Class Diagram', level=1)
doc.add_paragraph(
    'The following tables represent the system\'s domain classes, their attributes, '
    'methods, and relationships. This serves as a simplified class diagram.'
)
doc.add_paragraph()

classes = [
    ('User', [
        '+id: String (PK)',
        '+name: String',
        '+email: String (unique)',
        '+hashed_password: String',
        '+role: UserRole',
        '+avatar: String (nullable)',
        '+created_at: DateTime',
    ], ['+verify_password(plain): bool']),
    ('School', [
        '+id: String (PK)',
        '+name: String',
        '+code: String (unique)',
        '+description: Text',
        '+head_of_school: String (FK → User)',
        '+departments_count: Integer',
        '+status: Status',
        '+created_at: DateTime',
    ], []),
    ('Course', [
        '+id: String (PK)',
        '+title: String',
        '+description: Text',
        '+school_id: String (FK → School)',
        '+duration: Integer',
        '+duration_unit: String',
        '+fee: Float',
        '+pass_mark: Integer',
        '+status: Status',
        '+created_at: DateTime',
    ], ['+get_units(): List[CourseUnit]']),
    ('CourseUnit', [
        '+id: String (PK)',
        '+title: String',
        '+description: Text',
        '+course_id: String (FK → Course, nullable)',
        '+lecturer_id: String (FK → User, nullable)',
        '+credit_hours: Integer',
        '+status: ApprovalStatus',
        '+created_at: DateTime',
    ], ['+is_independent(): bool']),
    ('Intake', [
        '+id: String (PK)',
        '+name: String',
        '+course_id: String (FK → Course)',
        '+start_date: Date',
        '+end_date: Date',
        '+capacity: Integer',
        '+enrolled_count: Integer',
        '+fee_override: Float (nullable)',
        '+status: Status',
    ], ['+is_open(): bool']),
    ('Enrollment', [
        '+id: String (PK)',
        '+student_id: String (FK → User)',
        '+course_id: String (FK → Course)',
        '+intake_id: String (FK → Intake)',
        '+enrollment_date: Date',
        '+status: EnrollmentStatus',
        '+payment_status: PaymentStatus',
    ], ['+complete_payment()']),
    ('Payment', [
        '+id: String (PK)',
        '+transaction_id: String (unique)',
        '+student_id: String (FK → User)',
        '+amount: Float',
        '+date: DateTime',
        '+method: PaymentMethod',
        '+status: PaymentStatus',
    ], ['+retry()']),
    ('StudyMaterial', [
        '+id: String (PK)',
        '+title: String',
        '+course_unit_id: String (FK → CourseUnit)',
        '+type: MaterialType',
        '+file_url: String',
        '+file_size: Integer',
        '+uploaded_by: String (FK → User)',
    ], []),
    ('VirtualClass', [
        '+id: String (PK)',
        '+title: String',
        '+course_unit_id: String (FK → CourseUnit)',
        '+lecturer_id: String (FK → User)',
        '+date: Date',
        '+platform: Platform',
        '+meeting_link: String',
        '+is_live: Boolean',
    ], ['+get_join_url(): String']),
    ('Examination', [
        '+id: String (PK)',
        '+title: String',
        '+course_unit_id: String (FK → CourseUnit)',
        '+type: ExamType',
        '+pass_mark: Integer',
        '+time_limit: Integer (nullable)',
        '+max_attempts: Integer',
        '+status: Status',
    ], ['+is_available(): bool']),
    ('Certificate', [
        '+id: String (PK)',
        '+serial_number: String (unique)',
        '+student_id: String (FK → User)',
        '+course_id: String (FK → Course)',
        '+issue_date: Date',
        '+status: CertificateStatus',
    ], ['+generate_pdf(): bytes', '+verify(): bool']),
    ('TutoringSession', [
        '+id: String (PK)',
        '+tutor_id: String (FK → User)',
        '+student_id: String (FK → User)',
        '+subject: String',
        '+duration: Integer',
        '+hourly_rate: Float',
        '+status: SessionStatus',
    ], ['+get_total_cost(): float']),
    ('Notification', [
        '+id: String (PK)',
        '+user_id: String (FK → User)',
        '+title: String',
        '+message: String',
        '+category: NotificationCategory',
        '+is_read: Boolean',
        '+created_at: DateTime',
    ], ['+mark_as_read()']),
]

for class_name, attributes, methods in classes:
    # Class name heading
    p = doc.add_paragraph()
    run = p.add_run(class_name)
    run.bold = True
    run.font.size = Pt(12)

    # Create a table for the class
    row_count = 1 + len(attributes) + (1 if methods else 0) + len(methods)
    t = doc.add_table(rows=row_count, cols=1)
    t.style = 'Table Grid'

    # Header row
    t.rows[0].cells[0].text = f'<<class>> {class_name}'
    for para in t.rows[0].cells[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            r.bold = True

    # Attributes
    row_idx = 1
    for attr in attributes:
        t.rows[row_idx].cells[0].text = attr
        row_idx += 1

    # Methods separator + methods
    if methods:
        t.rows[row_idx].cells[0].text = '── Methods ──'
        for para in t.rows[row_idx].cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.italic = True
        row_idx += 1
        for method in methods:
            t.rows[row_idx].cells[0].text = method
            row_idx += 1

    doc.add_paragraph()  # spacing

doc.add_page_break()

# Section 7 - Relationships
doc.add_heading('7. Class Relationships', level=1)

relationships = [
    ('User', '1', 'School', '*', 'heads (as head_of_school)'),
    ('School', '1', 'Course', '*', 'contains'),
    ('Course', '1', 'CourseUnit', '*', 'has (optional — units can be independent)'),
    ('User', '1', 'CourseUnit', '*', 'teaches (as lecturer)'),
    ('Course', '1', 'Intake', '*', 'offered in'),
    ('User', '1', 'Enrollment', '*', 'enrolls in (as student)'),
    ('Course', '1', 'Enrollment', '*', 'enrolled for'),
    ('Intake', '1', 'Enrollment', '*', 'belongs to'),
    ('Enrollment', '1', 'Payment', '0..1', 'paid by'),
    ('CourseUnit', '1', 'StudyMaterial', '*', 'contains'),
    ('User', '1', 'StudyMaterial', '*', 'uploads'),
    ('CourseUnit', '1', 'VirtualClass', '*', 'scheduled for'),
    ('User', '1', 'VirtualClass', '*', 'conducts (as lecturer)'),
    ('CourseUnit', '1', 'Examination', '*', 'assessed by'),
    ('Examination', '1', 'ExamSubmission', '*', 'has submissions'),
    ('User', '1', 'Certificate', '*', 'earns (as student)'),
    ('Course', '1', 'Certificate', '*', 'certifies completion of'),
    ('User', '1', 'TutoringSession', '*', 'tutors / books'),
    ('User', '1', 'Notification', '*', 'receives'),
]

rel_table = doc.add_table(rows=len(relationships) + 1, cols=5)
rel_table.style = 'Table Grid'
rel_headers = ['From Class', 'Cardinality', 'To Class', 'Cardinality', 'Relationship']
for i, h in enumerate(rel_headers):
    rel_table.rows[0].cells[i].text = h
    for p in rel_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for idx, (from_cls, from_card, to_cls, to_card, rel) in enumerate(relationships, 1):
    rel_table.rows[idx].cells[0].text = from_cls
    rel_table.rows[idx].cells[1].text = from_card
    rel_table.rows[idx].cells[2].text = to_cls
    rel_table.rows[idx].cells[3].text = to_card
    rel_table.rows[idx].cells[4].text = rel

doc.add_paragraph()

# Section 8 - Enumerations
doc.add_heading('8. Enumerations', level=1)

enums = [
    ('UserRole', ['super_admin', 'admin', 'lecturer', 'student']),
    ('Status', ['active', 'inactive']),
    ('ApprovalStatus', ['active', 'pending_approval', 'rejected']),
    ('EnrollmentStatus', ['active', 'payment_pending', 'completed', 'dropped']),
    ('PaymentStatus', ['completed', 'pending', 'failed']),
    ('PaymentMethod', ['mobile_money', 'bank_transfer', 'card']),
    ('ExamType', ['quiz', 'assignment', 'final_exam']),
    ('MaterialType', ['pdf', 'video', 'presentation', 'assignment']),
    ('Platform', ['zoom', 'jitsi']),
    ('CertificateStatus', ['valid', 'revoked']),
    ('SessionStatus', ['upcoming', 'completed', 'cancelled']),
    ('NotificationCategory', ['enrollment', 'payment', 'class', 'exam']),
]

for enum_name, values in enums:
    p = doc.add_paragraph()
    run = p.add_run(f'{enum_name}: ')
    run.bold = True
    p.add_run(', '.join(values))

doc.add_paragraph()

# Section 9 - Module Dependencies
doc.add_heading('9. Module Dependency Summary', level=1)

dep_table = doc.add_table(rows=15, cols=3)
dep_table.style = 'Table Grid'
dep_headers = ['Module', 'Depends On', 'Provides To']
for i, h in enumerate(dep_headers):
    dep_table.rows[0].cells[i].text = h
    for p in dep_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

deps = [
    ('Authentication', 'Users DB', 'JWT tokens to all modules'),
    ('User Management', 'Authentication', 'User records to all modules'),
    ('School Management', 'Users', 'School IDs to Courses'),
    ('Course Management', 'Schools, Course Units', 'Course IDs to Intakes, Enrollment'),
    ('Course Units', 'Courses (optional), Users', 'Unit IDs to Materials, Classes, Exams'),
    ('Intake Management', 'Courses', 'Intake IDs to Enrollment'),
    ('Enrollment', 'Users, Courses, Intakes, Payments', 'Access control for content'),
    ('Study Materials', 'Course Units, Users', 'Learning content to Students'),
    ('Virtual Learning', 'Course Units, Users, Zoom/Jitsi', 'Live sessions'),
    ('Examination', 'Course Units, Users', 'Scores to Certification'),
    ('Certification', 'Courses, Users, Exams', 'Verifiable certificates'),
    ('Payment', 'Users, Interswitch', 'Payment status to Enrollment'),
    ('Tutoring', 'Users, Payments', 'Session records'),
    ('Notification', 'All modules (events)', 'Alerts to Users'),
]
for idx, (mod, dep, provides) in enumerate(deps, 1):
    dep_table.rows[idx].cells[0].text = mod
    dep_table.rows[idx].cells[1].text = dep
    dep_table.rows[idx].cells[2].text = provides

# Save
output = 'Makerere_Online_Architecture.docx'
doc.save(output)
print(f'Architecture document saved: {output}')
