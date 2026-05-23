"""Generate a Word document from SYSTEM_REPORT.md"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MAKERERE ONLINE UNIVERSITY\nMANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive System Functionality Report')
run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Document Version: 1.0\n').font.size = Pt(12)
info.add_run('Date: May 22, 2026\n').font.size = Pt(12)
info.add_run('Platform URL: https://makerereonlineschool.com\n').font.size = Pt(12)
info.add_run('API URL: https://api.makerereonlineschool.com').font.size = Pt(12)

doc.add_page_break()

# Read the markdown file
with open('SYSTEM_REPORT.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Parse and convert
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    # Skip the first title (we made our own title page)
    if line.startswith('# ') and i < 5:
        i += 1
        continue

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Skip table of contents header and links
    if line.strip().startswith('## Table of Contents'):
        i += 1
        while i < len(lines) and (lines[i].strip().startswith('1.') or
              lines[i].strip().startswith('2.') or
              lines[i].strip().startswith('3.') or
              lines[i].strip().startswith('4.') or
              lines[i].strip().startswith('5.') or
              lines[i].strip().startswith('6.') or
              lines[i].strip().startswith('7.') or
              lines[i].strip().startswith('8.') or
              lines[i].strip() == ''):
            i += 1
        continue

    # Heading 1 (##)
    if line.startswith('## '):
        text = line[3:].strip()
        # Remove markdown links
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        doc.add_heading(text, level=1)
        i += 1
        continue

    # Heading 2 (###)
    if line.startswith('### '):
        text = line[4:].strip()
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_heading(text, level=2)
        i += 1
        continue

    # Heading 3 (####)
    if line.startswith('#### '):
        text = line[5:].strip()
        doc.add_heading(text, level=3)
        i += 1
        continue

    # Table detection
    if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
        # Parse table
        headers = [c.strip() for c in line.split('|') if c.strip()]
        i += 2  # Skip header and separator
        rows = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip():
            cells = [c.strip() for c in lines[i].split('|') if c.strip()]
            rows.append(cells)
            i += 1

        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = re.sub(r'\*\*([^*]+)\*\*', r'\1', header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(headers):
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)

        doc.add_paragraph()
        continue

    # Bullet points
    if line.strip().startswith('- **'):
        # Bold label bullet
        match = re.match(r'\s*- \*\*([^*]+)\*\*:?\s*(.*)', line)
        if match:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(match.group(1) + ': ')
            run.bold = True
            p.add_run(match.group(2))
        i += 1
        continue

    if line.strip().startswith('- '):
        text = line.strip()[2:]
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue

    # Indented bullets
    if line.strip().startswith('  - '):
        text = line.strip()[4:]
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        p = doc.add_paragraph(text, style='List Bullet 2')
        i += 1
        continue

    # Regular paragraph
    if line.strip():
        text = line.strip()
        # Remove markdown formatting
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # Skip code blocks
        if text.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue

        doc.add_paragraph(text)

    i += 1

# Save
output_path = 'Makerere_Online_System_Report.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
